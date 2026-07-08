# -*- coding: utf-8 -*-
"""匯出：build_export 資料組裝 + Excel/Word 產檔並讀回驗證（重依賴 importorskip）。"""
import os
import sys
from datetime import date

import pytest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))

from cmuh_common.roster.export_common import (  # noqa: E402
    default_filename, leaves_summary, member_tally,
)
from cmuh_common.roster.service import RosterService  # noqa: E402
from cmuh_common.roster.storage import RosterStorage  # noqa: E402

YM = "2026-08"


def _svc(tmp_path):
    st = RosterStorage(str(tmp_path))
    st.save_config({
        "r_members": [{"id": "A", "name": "甲"}, {"id": "B", "name": "乙"}],
        "vs_members": [{"id": "D", "name": "D醫師"}],
        "points": {"weekday": 1, "weekend": 2, "national_holiday": 1},
        "duty_range_soft": [9, 11],
    })
    st.save_holiday_duty({"r": {date(2026, 8, 15): "A"}, "vs": {}})
    st.save_month(YM, {
        "r_duty": {"2026-08-01": {"person": "A"},
                   "2026-08-03": {"person": "B"}},
        "vs_duty": {"2026-08-01": {"person": "D"}},
        "leaves": {"r": {"A": ["2026-08-10"]}},
    })
    return RosterService(st)


# ─── 定案 PDF 留底 ───────────────────────────────────────────────────────────
def test_build_finalize_pdf_sections(tmp_path):
    svc = _svc(tmp_path)
    m = svc.storage.load_month(YM)
    m["report_r"] = "RR報告"
    m["report_vs"] = "VV報告"
    m["day_report"] = "DD日報告"
    svc.storage.save_month(YM, m)
    secs = svc.build_finalize_pdf_sections(YM)
    titles = [t for t, _ in secs]
    bodies = [b for _, b in secs]
    assert any("定案留底" in t for t in titles)          # 封面
    assert "RR報告" in bodies and "VV報告" in bodies and "DD日報告" in bodies


def test_archive_finalize_pdf_writes_pdf(tmp_path):
    pytest.importorskip("reportlab")
    svc = _svc(tmp_path)
    m = svc.storage.load_month(YM)
    m["report_r"] = "定案內容\n第二行"
    svc.storage.save_month(YM, m)
    path = svc.archive_finalize_pdf(YM)
    assert path.endswith("115年08月定案.pdf") and os.path.exists(path)
    with open(path, "rb") as f:
        assert f.read(4) == b"%PDF"                      # 真的是 PDF


def test_export_pdf_empty_sections_no_crash(tmp_path):
    pytest.importorskip("reportlab")
    from cmuh_common.roster import export_pdf
    out = tmp_path / "empty.pdf"
    export_pdf.export(str(out), [])
    assert out.exists() and out.read_bytes()[:4] == b"%PDF"


def test_export_pdf_wrap_splits_long_line():
    """codex(794124e)：過長行逐字斷行，不整段畫出頁面右緣（避免被裁）。"""
    from cmuh_common.roster.export_pdf import _wrap
    assert _wrap("abcdefghij", len, 4) == ["abcd", "efgh", "ij"]
    assert _wrap("ab", len, 4) == ["ab"]            # 不需斷
    assert _wrap("", len, 4) == [""]                # 空行保留間距


def test_export_pdf_long_line_produces_valid_pdf(tmp_path):
    pytest.importorskip("reportlab")
    from cmuh_common.roster import export_pdf
    out = tmp_path / "long.pdf"
    export_pdf.export(str(out), [("測試", "甲乙丙丁戊己庚辛" * 40)])   # 遠超一行寬
    assert out.exists() and out.read_bytes()[:4] == b"%PDF"


# ─── 純函式 ─────────────────────────────────────────────────────────────────
def test_build_export_structure(tmp_path):
    data = _svc(tmp_path).build_export(YM)
    assert data["year"] == 2026 and data["month"] == 8
    assert data["r"]["duty"][date(2026, 8, 1)] == "A"
    assert data["r"]["names"]["A"] == "甲"
    assert data["vs"]["names"]["D"] == "D"          # VS 用代號
    assert data["r"]["leaves"]["A"] == [date(2026, 8, 10)]
    assert date(2026, 8, 15) in data["holidays"]


def test_member_tally_and_summary(tmp_path):
    data = _svc(tmp_path).build_export(YM)
    tally = member_tally(data["r"], data["holidays"], data["params"])
    assert tally["A"] == {"wd": 0, "we": 1, "pt": 2}   # 8/1 週六
    assert tally["B"] == {"wd": 1, "we": 0, "pt": 1}   # 8/3 週一
    assert "甲(8/10)" in leaves_summary(data["r"])


def test_default_filename():
    data = {"year": 2026, "month": 7}
    assert default_filename(data, ".xlsx") == "115年07月班表.xlsx"


def test_member_tally_weekday_holiday_counts_as_holiday(tmp_path):
    """codex(33d8ecd)：平日的國定假日算假日班（we 欄），不進平日欄（wd）。"""
    svc = _svc(tmp_path)
    svc.storage.save_holiday_duty({"r": {date(2026, 8, 3): "X"}, "vs": {}})  # 8/3 週一設假日
    data = svc.build_export(YM)
    tally = member_tally(data["r"], data["holidays"], data["params"])
    assert tally["B"] == {"wd": 0, "we": 1, "pt": 1}   # 8/3 週一國定假日→假日班


def test_rf11_member_tally_includes_removed_member(tmp_path):
    """RF-11：值班者已不在 config 名單 → member_tally 仍動態納入、不漏點數。"""
    svc = _svc(tmp_path)
    cfg = svc.storage.load_config()
    cfg["r_members"] = [{"id": "A", "name": "甲"}]       # 移除 B（月檔 duty 仍有 B）
    svc.storage.save_config(cfg)
    data = svc.build_export(YM)
    tally = member_tally(data["r"], data["holidays"], data["params"])
    assert tally["B"] == {"wd": 1, "we": 0, "pt": 1}     # 8/3 週一


def test_rf11_summary_sheet_includes_removed_member(tmp_path):
    """RF-11：結算 sheet 補列已離名單的值班者（標「(已離)」、帳本印「—」）。"""
    pytest.importorskip("openpyxl")
    from openpyxl import load_workbook

    from cmuh_common.roster import export_xlsx
    svc = _svc(tmp_path)
    cfg = svc.storage.load_config()
    cfg["r_members"] = [{"id": "A", "name": "甲"}]
    svc.storage.save_config(cfg)
    data = svc.build_export(YM)
    out = tmp_path / "out.xlsx"
    export_xlsx.export(str(out), data)
    wb = load_workbook(str(out))
    rows = [[c.value for c in row] for row in wb["結算"].iter_rows(min_row=2)]
    b_rows = [r for r in rows if r[1] and str(r[1]).startswith("B")]
    assert b_rows, "結算 sheet 應補上已離名單的 B 一列"
    assert "已離" in b_rows[0][1]
    assert b_rows[0][2] == 1 and b_rows[0][5] == 1       # 平日班=1、點數=1
    assert b_rows[0][6] == "—"                           # 帳本作廢印 —


# ─── Excel 產檔讀回 ─────────────────────────────────────────────────────────
def test_export_xlsx_roundtrip(tmp_path):
    pytest.importorskip("openpyxl")
    from openpyxl import load_workbook

    from cmuh_common.roster import export_xlsx
    data = _svc(tmp_path).build_export(YM)
    out = tmp_path / "out.xlsx"
    export_xlsx.export(str(out), data)

    wb = load_workbook(str(out))
    assert "值班表" in wb.sheetnames and "結算" in wb.sheetnames
    ws = wb["值班表"]
    assert "115年08月" in ws["A1"].value
    # 8/1=週六 → 第一週列(row3) 週六欄(F)
    f3 = ws["F3"].value
    assert "R:甲" in f3 and "VS:D" in f3
    # 結算：甲 假日 1 班
    summ = wb["結算"]
    rows = [[c.value for c in row] for row in summ.iter_rows(min_row=2)]
    a_row = [r for r in rows if r[1] == "甲"][0]
    assert a_row[3] == 1 and a_row[5] == 2            # 假日欄=1、點數=2


# ─── Word 產檔讀回 ─────────────────────────────────────────────────────────
def test_export_docx_roundtrip(tmp_path):
    pytest.importorskip("docx")
    from docx import Document

    from cmuh_common.roster import export_docx
    data = _svc(tmp_path).build_export(YM)
    out = tmp_path / "out.docx"
    export_docx.export(str(out), data)

    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "115年08月" in text
    assert "醫師請假：" in text and "甲(8/10)" in text
    # 第一週表格：一線(row2)/三線(row3) 的週六欄(col6)
    t0 = doc.tables[0]
    assert t0.cell(2, 6).text == "甲"                # 一線=R
    assert t0.cell(3, 6).text == "D"                 # 三線=VS 代號


def test_export_empty_month_no_crash(tmp_path):
    """空月份（無排班）也能產檔不炸。"""
    pytest.importorskip("openpyxl")
    from cmuh_common.roster import export_xlsx
    st = RosterStorage(str(tmp_path))
    st.save_config({"r_members": [], "vs_members": []})
    data = RosterService(st).build_export(YM)
    export_xlsx.export(str(tmp_path / "empty.xlsx"), data)
    assert (tmp_path / "empty.xlsx").exists()


# ─── RP3-05：定案 PDF 符號淨化 ──────────────────────────────────────────────
def test_rp3_05_pdf_sanitize_maps_and_boxes():
    """[RP3-05] 已知符號換可讀替代；BMP 外字元一律 □（避免 MSung 字型畫成空框）。"""
    from cmuh_common.roster.export_pdf import _sanitize
    assert _sanitize("\U0001F512鎖定") == "[鎖]鎖定"
    assert _sanitize("✓通過 ✗失敗 ⚠注意") == "[v]通過 [x]失敗 [!]注意"
    assert _sanitize("純中文與 ASCII abc 不變") == "純中文與 ASCII abc 不變"
    assert _sanitize("\U0001F600") == "□"          # 未列入對照的 emoji → □
    assert _sanitize("・═") == ".="


def test_rp3_05_pdf_export_with_symbols_smoke(tmp_path):
    """含 🔒/⚠/emoji 的 sections 仍能產出 PDF（不因字型缺字崩）。"""
    pytest.importorskip("reportlab")
    from cmuh_common.roster import export_pdf
    out = str(tmp_path / "final.pdf")
    export_pdf.export(out, [("標題\U0001F512", "第一行 ✓\n第二行 ⚠ \U0001F600")])
    assert os.path.exists(out)


# ─── RS-01：PGY/Clerk 日排班匯出 ────────────────────────────────────────────
def test_rs01_fmt_cell_and_grid_rows():
    """[RS-01] 格字串順序＝照光→治療室→房→切片→放假；週三下午照光有人無治療室。"""
    from cmuh_common.roster.export_common import _fmt_day_cell, day_grid_rows
    assert _fmt_day_cell(
        {"照光": ["A"], "治療室": ["B"], "101": ["C"], "切片室": ["1"]}
    ) == "照光:A 治療室:B 101:C 切片室:1"
    assert _fmt_day_cell({"照光": ["A", "B"]}) == "照光:A、B"
    assert _fmt_day_cell({}) == ""

    day_slots = {
        "2026-08-03": {"上午": {"照光": ["A"], "治療室": ["B"], "101": ["C"]}},
        "2026-08-05": {"下午": {"照光": ["A"]}},          # 週三下午：只有照光
    }
    blocks = day_grid_rows(day_slots, 2026, 8)
    assert blocks
    # 找含 8/5（週三）的週，確認下午格＝照光有人、無治療室
    wed_cells = []
    for blk in blocks:
        for i, d in enumerate(blk["weekdays"]):
            if d and d.isoformat() == "2026-08-05":
                pm = dict(blk["sessions"])["下午"]
                wed_cells.append(pm[i])
    assert wed_cells == ["照光:A"]
    assert all("治療室" not in c for c in wed_cells)


def _day_export_data(tmp_path):
    st = RosterStorage(str(tmp_path))
    st.save_config({"r_members": [], "vs_members": []})
    m = st.load_month(YM)
    m["day_slots"] = {
        "2026-08-03": {"上午": {"照光": ["A"], "治療室": ["B"], "101": ["C"]}},
        "2026-08-05": {"下午": {"照光": ["A"]}},
    }
    st.save_month(YM, m)
    return RosterService(st).build_export(YM)


def test_rs01_xlsx_has_pgy_clerk_sheet(tmp_path):
    op = pytest.importorskip("openpyxl")
    from cmuh_common.roster import export_xlsx
    data = _day_export_data(tmp_path)
    path = str(tmp_path / "day.xlsx")
    export_xlsx.export(path, data)
    wb = op.load_workbook(path)
    assert "PGY-Clerk" in wb.sheetnames
    # 內容有出現照光/治療室（至少一格非空）
    ws = wb["PGY-Clerk"]
    joined = "".join(str(c.value or "") for row in ws.iter_rows() for c in row)
    assert "照光:A" in joined and "治療室:B" in joined


def test_rs01_docx_has_day_schedule(tmp_path):
    pytest.importorskip("docx")
    from docx import Document

    from cmuh_common.roster import export_docx
    data = _day_export_data(tmp_path)
    path = str(tmp_path / "day.docx")
    export_docx.export(path, data)
    doc = Document(path)
    heads = [p.text for p in doc.paragraphs]
    assert any("PGY / Clerk 日排班" in t for t in heads)
