# -*- coding: utf-8 -*-
"""匯出 Word 值班表（python-docx，重依賴 → 呼叫端負責 lazy 安裝）。

仿 115-07 月班表：頁首標題 + 「醫師請假」摘要 + 每週一個 4 列×8 欄表格
（列＝日期／星期／一線(R)／三線(VS)；欄1＝標籤，欄2-8＝週一..週日）。
"""
from __future__ import annotations

from cmuh_common.roster.export_common import (
    WD_CN, day_grid_rows, leaves_summary, title_text,
)
from cmuh_common.roster.model import week_matrix


def export(path: str, data: dict) -> None:
    from docx import Document  # noqa: PLC0415 (lazy 重依賴)
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    h = doc.add_heading(title_text(data), level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 醫師請假（R + VS 合併一段）
    r_leave = leaves_summary(data["r"])
    vs_leave = leaves_summary(data["vs"])
    merged = "　".join(s for s in (r_leave, vs_leave) if s)
    doc.add_paragraph(f"醫師請假：{merged or '（無）'}")

    r, vs = data["r"], data["vs"]
    for week in week_matrix(data["year"], data["month"]):
        table = doc.add_table(rows=4, cols=8)
        table.style = "Table Grid"
        labels = ("日期", "星期", "一線", "三線")
        for i, lab in enumerate(labels):
            table.cell(i, 0).text = lab
        for c, d in enumerate(week, start=1):
            if d is None:
                continue
            table.cell(0, c).text = str(d.day)
            table.cell(1, c).text = WD_CN[d.weekday()]
            rp = r["duty"].get(d)
            vp = vs["duty"].get(d)
            rtext = r["names"].get(rp, rp) if rp else ""
            bp = (data.get("saturday_biopsy") or {}).get(d)   # [週六切片]
            if bp:
                rtext = (rtext + "\n" if rtext else "") + \
                    f"切:{r['names'].get(bp, bp)}"
            table.cell(2, c).text = rtext
            table.cell(3, c).text = vs["names"].get(vp, vp) if vp else ""
        doc.add_paragraph("")            # 週表格之間留白

    _add_day_schedule(doc, data)         # [RS-01] PGY/Clerk 日排班
    doc.save(path)


def _add_day_schedule(doc, data: dict) -> None:
    """[RS-01] 附上 PGY/Clerk 週格網：每週一表格（首列＝日期；其後上午/下午各一列，
    欄1＝時段標籤、欄2-6＝週一~五）。無日排班則整段略過。"""
    blocks = day_grid_rows(data.get("day_slots") or {}, data["year"], data["month"])
    if not blocks:
        return
    doc.add_heading("PGY / Clerk 日排班", level=2)
    for blk in blocks:
        rows = 1 + len(blk["sessions"])
        table = doc.add_table(rows=rows, cols=6)
        table.style = "Table Grid"
        table.cell(0, 0).text = "日期"
        for c, d in enumerate(blk["weekdays"], start=1):
            table.cell(0, c).text = (f"{d.month}/{d.day}（{WD_CN[d.weekday()]}）"
                                     if d else "")
        for r, (sess, cells) in enumerate(blk["sessions"], start=1):
            table.cell(r, 0).text = sess
            for c, val in enumerate(cells, start=1):
                table.cell(r, c).text = val
        doc.add_paragraph("")
